#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
UltimateFileConverter - 终极文件格式转换器
功能：图片、音视频、文档等多格式转换、批量处理、预览
作者：LYP
GitHub：https://github.com/lyp0746
邮箱：1610369302@qq.com
版本：4.0.0
"""

import sys
import os
import io
import time
import base64
import hashlib
import shutil
import zipfile
import tarfile
import threading
import subprocess
import configparser
import json
from pathlib import Path
from typing import Optional, List, Dict, Any
from datetime import datetime

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QFileDialog, QMessageBox, QProgressBar,
    QTextEdit, QComboBox, QSlider, QCheckBox, QTabWidget,
    QGroupBox, QRadioButton, QButtonGroup, QSplitter, QListWidget,
    QTableWidget, QTableWidgetItem, QStatusBar, QMenuBar, QMenu,
    QAction, QToolBar, QDialog, QDialogButtonBox, QSpinBox,
    QLineEdit, QTreeWidget, QTreeWidgetItem, QFrame, QScrollArea,
    QGridLayout
)
from PyQt5.QtCore import (
    Qt, QThread, pyqtSignal, QSize, QTimer, QSettings,
    QMimeData, QUrl, QEvent
)
from PyQt5.QtGui import (
    QIcon, QFont, QColor, QPalette, QDragEnterEvent,
    QDropEvent, QPixmap, QImage, QPainter, QLinearGradient
)

# ================== 依赖库检测 ==================
try:
    from PIL import Image, ImageDraw, ImageFont, ImageFilter, ImageEnhance, ExifTags

    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import pandas as pd

    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    from docx import Document
    from bs4 import BeautifulSoup

    DOC_AVAILABLE = True
except ImportError:
    DOC_AVAILABLE = False

try:
    from pypdf import PdfReader
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.units import inch

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from pydub import AudioSegment

    AUDIO_AVAILABLE = True
except ImportError:
    AUDIO_AVAILABLE = False

try:
    from moviepy.video.io.VideoFileClip import VideoFileClip

    VIDEO_AVAILABLE = True
except ImportError:
    VIDEO_AVAILABLE = False

try:
    import ebooklib
    from ebooklib import epub

    EPUB_AVAILABLE = True
except ImportError:
    EPUB_AVAILABLE = False

try:
    import yaml

    YAML_AVAILABLE = True
except ImportError:
    YAML_AVAILABLE = False

try:
    import qrcode

    QRCODE_AVAILABLE = True
except ImportError:
    QRCODE_AVAILABLE = False

try:
    import cairosvg

    CAIROSVG_AVAILABLE = True
except ImportError:
    CAIROSVG_AVAILABLE = False

try:
    import toml

    TOML_AVAILABLE = True
except ImportError:
    TOML_AVAILABLE = False

CALIBRE_AVAILABLE = shutil.which("ebook-convert") is not None


# ================== 转换工作线程 ==================
class ConversionWorker(QThread):
    """转换工作线程"""
    progress = pyqtSignal(int, int, str)  # current, total, message
    finished = pyqtSignal(int, int, float)  # success, failed, elapsed
    log = pyqtSignal(str, str)  # message, level

    def __init__(self, converter, files, settings):
        super().__init__()
        self.converter = converter
        self.files = files
        self.settings = settings
        self.is_running = True

    def run(self):
        start_time = time.time()
        success = 0
        failed = 0
        total = len(self.files)

        for i, input_file in enumerate(self.files):
            if not self.is_running:
                break

            try:
                filename = Path(input_file).name
                self.log.emit(f"[{i + 1}/{total}] 转换: {filename}", "info")
                self.progress.emit(i + 1, total, f"正在处理: {filename}")

                output_file = self.converter.convert_single_file(
                    input_file, self.settings
                )

                if output_file and Path(output_file).exists():
                    success += 1
                    self.log.emit(f"✅ 成功: {Path(output_file).name}", "success")
                else:
                    failed += 1
                    self.log.emit("❌ 失败: 输出文件未生成", "error")

            except Exception as e:
                failed += 1
                self.log.emit(f"❌ 失败: {str(e)}", "error")

        elapsed = time.time() - start_time
        self.finished.emit(success, failed, elapsed)

    def stop(self):
        self.is_running = False


# ================== 转换引擎 ==================
class ConversionEngine:
    """转换引擎核心类"""

    def __init__(self):
        self.chinese_font_registered = False
        self._register_chinese_fonts()

    def _register_chinese_fonts(self):
        """注册中文字体"""
        if not PDF_AVAILABLE:
            return

        try:
            font_paths = []
            if sys.platform == "win32":
                font_paths = [
                    "C:/Windows/Fonts/simsun.ttc",
                    "C:/Windows/Fonts/msyh.ttc",
                    "C:/Windows/Fonts/simhei.ttf",
                ]
            elif sys.platform == "darwin":
                font_paths = [
                    "/System/Library/Fonts/PingFang.ttc",
                    "/Library/Fonts/Arial Unicode.ttf",
                ]
            else:
                font_paths = [
                    "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
                    "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
                    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                ]

            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont("ChineseFont", font_path))
                        self.chinese_font_registered = True
                        break
                    except:
                        continue
        except:
            pass

    def convert_single_file(self, input_file: str, settings: Dict[str, Any]) -> Optional[str]:
        """转换单个文件"""
        input_path = Path(input_file)
        output_format = settings['output_format'].lower()
        category = settings['category']

        # 确定输出路径
        if settings.get('output_folder') and os.path.isdir(settings['output_folder']):
            output_path = Path(settings['output_folder']) / f"{input_path.stem}.{output_format}"
        else:
            output_path = input_path.with_suffix(f".{output_format}")

        if output_path == input_path:
            output_path = input_path.with_suffix(f".converted.{output_format}")

        # 根据类别调用相应的转换函数
        try:
            if "图片" in category:
                result = self.convert_image(input_path, output_path, settings)
            elif "文档" in category:
                result = self.convert_document(input_path, output_path, settings)
            elif "数据" in category:
                result = self.convert_data(input_path, output_path, settings)
            elif "音频" in category:
                result = self.convert_audio(input_path, output_path, settings)
            elif "视频" in category:
                result = self.convert_video(input_path, output_path, settings)
            elif "压缩" in category:
                result = self.convert_archive(input_path, output_path, settings)
            elif "电子书" in category:
                result = self.convert_ebook(input_path, output_path, settings)
            elif "编码" in category:
                result = self.convert_encoding(input_path, output_path, settings)
            elif "特殊" in category:
                result = self.special_functions(input_path, output_path, settings)
            else:
                raise Exception(f"不支持的类别: {category}")

            return str(result) if result else None

        except Exception as e:
            raise Exception(f"转换失败: {str(e)}")

    # ========== 图片转换 ==========
    def convert_image(self, input_path: Path, output_path: Path, settings: Dict) -> Optional[Path]:
        """图片转换（增强版）"""
        if not PIL_AVAILABLE:
            raise Exception("需要安装 Pillow: pip install Pillow")

        try:
            input_ext = input_path.suffix.lower()[1:]
            output_ext = output_path.suffix.lower()[1:]
            quality = settings.get('quality', 85)

            # SVG 特殊处理
            if input_ext == "svg":
                if output_ext == "svg":
                    shutil.copy2(input_path, output_path)
                    return output_path
                elif CAIROSVG_AVAILABLE:
                    if output_ext == "pdf":
                        cairosvg.svg2pdf(url=str(input_path), write_to=str(output_path))
                    elif output_ext in ["png", "jpg", "jpeg"]:
                        temp_png = output_path.with_suffix(".png")
                        cairosvg.svg2png(url=str(input_path), write_to=str(temp_png))
                        if output_ext in ["jpg", "jpeg"]:
                            img = Image.open(temp_png)
                            img_rgb = img.convert("RGB")
                            img_rgb.save(output_path, quality=quality, optimize=True)
                            temp_png.unlink()
                        else:
                            return temp_png
                    else:
                        temp_png = input_path.with_suffix(".temp.png")
                        cairosvg.svg2png(url=str(input_path), write_to=str(temp_png))
                        img = Image.open(temp_png)
                        img.save(output_path)
                        temp_png.unlink()
                    return output_path
                else:
                    raise Exception("SVG转换需要安装: pip install cairosvg")

            # 标准图片处理
            img = Image.open(input_path)

            # EXIF 方向处理
            try:
                exif = img._getexif() if hasattr(img, "_getexif") else None
                if exif:
                    for k, v in ExifTags.TAGS.items():
                        if v == "Orientation":
                            orientation = exif.get(k)
                            if orientation == 3:
                                img = img.rotate(180, expand=True)
                            elif orientation == 6:
                                img = img.rotate(270, expand=True)
                            elif orientation == 8:
                                img = img.rotate(90, expand=True)
                            break
            except:
                pass

            # PDF 输出
            if output_ext == "pdf":
                if not PDF_AVAILABLE:
                    raise Exception("需要安装 reportlab")
                img_rgb = img.convert("RGB")
                img_rgb.save(output_path, "PDF", resolution=100.0)
                return output_path

            # SVG 输出不支持
            if output_ext == "svg":
                raise Exception("位图转SVG需要专业矢量化工具")

            # RGBA 到 RGB
            if img.mode in ("RGBA", "LA", "P") and output_ext in ["jpg", "jpeg", "bmp"]:
                bg = Image.new("RGB", img.size, (255, 255, 255))
                if img.mode == "P":
                    img = img.convert("RGBA")
                if img.mode in ("RGBA", "LA"):
                    bg.paste(img, mask=img.split()[-1])
                else:
                    bg.paste(img)
                img = bg

            # ICO 特殊处理
            if output_ext == "ico":
                img.save(output_path, format="ICO", sizes=[(32, 32), (64, 64), (128, 128)])
                return output_path

            # 保存参数
            save_kwargs = {}
            if output_ext in ["jpg", "jpeg"]:
                save_kwargs = {"quality": quality, "optimize": True, "progressive": True}
            elif output_ext == "png":
                save_kwargs = {"optimize": True}
            elif output_ext == "webp":
                save_kwargs = {"quality": quality, "method": 6}

            img.save(output_path, **save_kwargs)
            return output_path

        except Exception as e:
            raise Exception(f"图片转换失败: {str(e)}")

    # ========== 文档转换 ==========
    def convert_document(self, input_path: Path, output_path: Path, settings: Dict) -> Optional[Path]:
        """文档转换（支持图片）"""
        try:
            input_ext = input_path.suffix.lower()[1:]
            output_ext = output_path.suffix.lower()[1:]

            content = ""
            images = []

            # 读取
            if input_ext == "docx" and DOC_AVAILABLE:
                doc = Document(input_path)
                for para in doc.paragraphs:
                    if para.text.strip():
                        content += para.text + "\n\n"

            elif input_ext == "pdf" and PDF_AVAILABLE:
                reader = PdfReader(str(input_path))
                texts = [page.extract_text() or "" for page in reader.pages]
                content = "\n\n".join(texts)

            elif input_ext == "html":
                with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
                    html_content = f.read()
                if DOC_AVAILABLE:
                    soup = BeautifulSoup(html_content, "html.parser")
                    content = soup.get_text()
                else:
                    content = html_content
            else:
                with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()

            # 写入
            if output_ext == "html":
                html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{input_path.stem}</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            padding: 20px;
            max-width: 900px;
            margin: 0 auto;
        }}
        pre {{ white-space: pre-wrap; word-wrap: break-word; }}
    </style>
</head>
<body>
    <h1>{input_path.stem}</h1>
    <pre>{content}</pre>
</body>
</html>"""
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(html_content)

            elif output_ext == "pdf" and PDF_AVAILABLE:
                doc = SimpleDocTemplate(str(output_path), pagesize=letter)
                styles = getSampleStyleSheet()

                if self.chinese_font_registered:
                    style = ParagraphStyle(
                        "ChineseStyle",
                        parent=styles["Normal"],
                        fontName="ChineseFont",
                        fontSize=12,
                        leading=18,
                    )
                else:
                    style = styles["Normal"]

                story = []
                for para in content.split("\n\n"):
                    if para.strip():
                        safe_para = para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                        safe_para = safe_para.replace("\n", "<br/>")
                        try:
                            story.append(Paragraph(safe_para, style))
                            story.append(Spacer(1, 12))
                        except:
                            pass
                doc.build(story)

            elif output_ext == "docx" and DOC_AVAILABLE:
                new_doc = Document()
                for para in content.split("\n\n"):
                    if para.strip():
                        new_doc.add_paragraph(para)
                new_doc.save(str(output_path))
            else:
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(content)

            return output_path

        except Exception as e:
            raise Exception(f"文档转换失败: {str(e)}")

    # ========== 数据转换 ==========
    def convert_data(self, input_path: Path, output_path: Path, settings: Dict) -> Optional[Path]:
        """数据格式转换"""
        if not PANDAS_AVAILABLE:
            raise Exception("需要安装 pandas: pip install pandas openpyxl")

        try:
            input_ext = input_path.suffix.lower()[1:]
            output_ext = output_path.suffix.lower()[1:]

            # 读取
            if input_ext == "csv":
                df = pd.read_csv(input_path, encoding="utf-8")
            elif input_ext in ["xlsx", "xls"]:
                df = pd.read_excel(input_path)
            elif input_ext == "json":
                df = pd.read_json(input_path)
            elif input_ext == "xml":
                df = pd.read_xml(input_path)
            elif input_ext in ["yaml", "yml"] and YAML_AVAILABLE:
                with open(input_path, "r", encoding="utf-8") as f:
                    data = yaml.safe_load(f)
                df = pd.DataFrame(data)
            elif input_ext == "parquet":
                df = pd.read_parquet(input_path)
            else:
                df = pd.read_csv(input_path, encoding="utf-8")

            # 写入
            if output_ext == "csv":
                df.to_csv(output_path, index=False, encoding="utf-8-sig")
            elif output_ext == "xlsx":
                df.to_excel(output_path, index=False, engine="openpyxl")
            elif output_ext == "json":
                df.to_json(output_path, orient="records", force_ascii=False, indent=2)
            elif output_ext == "xml":
                df.to_xml(output_path, index=False)
            elif output_ext in ["yaml", "yml"] and YAML_AVAILABLE:
                data = df.to_dict(orient="records")
                with open(output_path, "w", encoding="utf-8") as f:
                    yaml.dump(data, f, allow_unicode=True, default_flow_style=False)
            elif output_ext == "html":
                html = df.to_html(index=False, border=1)
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(f"<!DOCTYPE html><html><head><meta charset='UTF-8'></head><body>{html}</body></html>")
            else:
                raise Exception(f"不支持的输出格式: {output_ext}")

            return output_path

        except Exception as e:
            raise Exception(f"数据转换失败: {str(e)}")

    # ========== 音频转换 ==========
    def convert_audio(self, input_path: Path, output_path: Path, settings: Dict) -> Optional[Path]:
        """音频格式转换"""
        if not AUDIO_AVAILABLE:
            raise Exception("需要安装 pydub + FFmpeg")

        try:
            input_ext = input_path.suffix.lower()[1:]
            output_ext = output_path.suffix.lower()[1:]
            quality = settings.get('quality', 85)

            audio = AudioSegment.from_file(str(input_path), format=input_ext)

            export_params = {"format": output_ext}
            if output_ext in ["mp3", "ogg", "wma"]:
                bitrate = f"{max(32, int(quality * 3.2))}k"
                export_params["bitrate"] = bitrate

            audio.export(str(output_path), **export_params)
            return output_path

        except Exception as e:
            raise Exception(f"音频转换失败: {str(e)}")

    # ========== 视频转换 ==========
    def convert_video(self, input_path: Path, output_path: Path, settings: Dict) -> Optional[Path]:
        """视频格式转换"""
        if not VIDEO_AVAILABLE:
            raise Exception("需要安装 moviepy + FFmpeg")

        try:
            output_ext = output_path.suffix.lower()[1:]
            quality = settings.get('quality', 85)

            with VideoFileClip(str(input_path)) as clip:
                if output_ext == "gif":
                    clip_resized = clip.resize(width=480)
                    clip_resized.write_gif(str(output_path), fps=10, program="ffmpeg")
                else:
                    codec_map = {
                        "mp4": "libx264", "avi": "mpeg4", "mkv": "libx264",
                        "mov": "libx264", "flv": "flv", "wmv": "wmv2",
                        "webm": "libvpx-vp9", "m4v": "libx264", "3gp": "libx264",
                    }
                    codec = codec_map.get(output_ext, "libx264")
                    bitrate = f"{max(300, int(quality * 50))}k"

                    clip.write_videofile(
                        str(output_path),
                        codec=codec,
                        bitrate=bitrate,
                        audio_codec="aac",
                        logger=None
                    )

            return output_path

        except Exception as e:
            raise Exception(f"视频转换失败: {str(e)}")

    # ========== 压缩文件转换 ==========
    def convert_archive(self, input_path: Path, output_path: Path, settings: Dict) -> Optional[Path]:
        """压缩格式转换"""
        try:
            output_ext = output_path.suffix.lower()[1:]

            temp_dir = None
            if input_path.suffix.lower() == ".zip":
                temp_dir = input_path.parent / f"temp_{input_path.stem}"
                temp_dir.mkdir(exist_ok=True)
                with zipfile.ZipFile(input_path, "r") as zip_ref:
                    zip_ref.extractall(temp_dir)
                source_path = temp_dir
            elif ".tar" in "".join(input_path.suffixes):
                temp_dir = input_path.parent / f"temp_{input_path.stem}"
                temp_dir.mkdir(exist_ok=True)
                with tarfile.open(input_path, "r:*") as tar_ref:
                    tar_ref.extractall(temp_dir)
                source_path = temp_dir
            else:
                source_path = input_path if input_path.is_dir() else input_path.parent

            # 创建压缩包
            if output_ext == "zip":
                with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zipf:
                    for file_path in Path(source_path).rglob("*"):
                        if file_path.is_file():
                            zipf.write(file_path, file_path.relative_to(source_path))
            elif ".tar" in str(output_path):
                mode = "w:gz" if output_path.suffix == ".gz" else "w:bz2" if output_path.suffix == ".bz2" else "w"
                with tarfile.open(output_path, mode) as tarf:
                    tarf.add(source_path, arcname=source_path.name)
            else:
                raise Exception("仅支持 zip / tar / tar.gz / tar.bz2")

            if temp_dir and temp_dir.exists():
                shutil.rmtree(temp_dir)

            return output_path

        except Exception as e:
            raise Exception(f"压缩文件转换失败: {str(e)}")

    # ========== 电子书转换 ==========
    def convert_ebook(self, input_path: Path, output_path: Path, settings: Dict) -> Optional[Path]:
        """电子书格式转换"""
        try:
            input_ext = input_path.suffix.lower()[1:]
            output_ext = output_path.suffix.lower()[1:]

            # 优先使用 Calibre
            if CALIBRE_AVAILABLE and input_ext in ("epub", "mobi", "azw3"):
                try:
                    result = subprocess.run(
                        ["ebook-convert", str(input_path), str(output_path)],
                        stdout=subprocess.PIPE,
                        stderr=subprocess.PIPE,
                        text=True
                    )
                    if result.returncode == 0:
                        return output_path
                except:
                    pass

            # 文本提取
            content = ""
            if input_ext == "epub" and EPUB_AVAILABLE:
                book = epub.read_epub(str(input_path))
                for item in book.get_items():
                    if item.get_type() == ebooklib.ITEM_DOCUMENT:
                        try:
                            html_content = item.get_content().decode("utf-8", errors="ignore")
                            if DOC_AVAILABLE:
                                soup = BeautifulSoup(html_content, "html.parser")
                                content += soup.get_text() + "\n\n"
                            else:
                                content += html_content + "\n\n"
                        except:
                            pass
            elif input_ext == "pdf" and PDF_AVAILABLE:
                reader = PdfReader(str(input_path))
                texts = [page.extract_text() or "" for page in reader.pages]
                content = "\n\n".join(texts)
            else:
                with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()

            # 输出
            if output_ext == "txt":
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(content)
            elif output_ext == "html":
                html = f"""<!DOCTYPE html>
<html>
<head><meta charset="UTF-8"><title>{input_path.stem}</title></head>
<body><h1>{input_path.stem}</h1><pre>{content}</pre></body>
</html>"""
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(html)
            else:
                raise Exception(f"电子书转换到 {output_ext} 需要安装 Calibre")

            return output_path

        except Exception as e:
            raise Exception(f"电子书转换失败: {str(e)}")

    # ========== 编码转换 ==========
    def convert_encoding(self, input_path: Path, output_path: Path, settings: Dict) -> Optional[Path]:
        """编码转换"""
        try:
            with open(input_path, "rb") as f:
                data = f.read()

            output_format = settings['output_format'].lower()

            if output_format == "base64":
                encoded = base64.b64encode(data).decode("ascii")
            elif output_format == "hex":
                encoded = data.hex()
            elif output_format == "md5":
                encoded = hashlib.md5(data).hexdigest()
            elif output_format == "sha256":
                encoded = hashlib.sha256(data).hexdigest()
            else:
                encoded = str(data)

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(encoded)

            return output_path

        except Exception as e:
            raise Exception(f"编码转换失败: {str(e)}")

    # ========== 特殊功能 ==========
    def special_functions(self, input_path: Path, output_path: Path, settings: Dict) -> Optional[Path]:
        """特殊功能（二维码、缩略图等）"""
        if not PIL_AVAILABLE:
            raise Exception("需要安装 Pillow")

        try:
            func = settings['output_format'].lower()
            quality = settings.get('quality', 85)

            if func == "qrcode":
                if not QRCODE_AVAILABLE:
                    raise Exception("需要安装 qrcode")

                with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
                    data = f.read()

                qr = qrcode.QRCode(
                    version=1,
                    error_correction=qrcode.constants.ERROR_CORRECT_H,
                    box_size=10,
                    border=4,
                )
                qr.add_data(data)
                qr.make(fit=True)

                img = qr.make_image(fill_color="black", back_color="white")
                final_output = output_path.with_suffix(".png")
                img.save(final_output, format="PNG")
                return final_output

            elif func == "thumbnail":
                img = Image.open(input_path)
                img.thumbnail((200, 200), Image.Resampling.LANCZOS)
                final_output = output_path.with_suffix(".png")
                img.save(final_output, quality=quality)
                return final_output

            elif func == "watermark":
                img = Image.open(input_path).convert("RGBA")
                watermark = Image.new("RGBA", img.size, (0, 0, 0, 0))
                draw = ImageDraw.Draw(watermark)

                text = "CONVERTED"
                try:
                    font = ImageFont.truetype("arial.ttf", 48)
                except:
                    font = ImageFont.load_default()

                bbox = draw.textbbox((0, 0), text, font=font)
                position = ((img.size[0] - bbox[2]) // 2, (img.size[1] - bbox[3]) // 2)
                draw.text(position, text, fill=(255, 255, 255, 128), font=font)

                result = Image.alpha_composite(img, watermark)
                final_output = output_path.with_suffix(".png")
                result.convert("RGB").save(final_output, quality=quality)
                return final_output

            else:
                raise Exception(f"不支持的特殊功能: {func}")

        except Exception as e:
            raise Exception(f"特殊功能失败: {str(e)}")


# ================== 主窗口 ==================
class UltimateFileConverter(QMainWindow):
    """主窗口类"""

    def __init__(self):
        super().__init__()

        # 初始化数据
        self.input_files = []
        self.output_folder = None
        self.conversion_history = []
        self.engine = ConversionEngine()
        self.current_theme = "light"

        # 设置管理
        self.settings = QSettings("UltimateConverter", "Settings")

        # 格式类别定义
        self.format_categories = {
            "🖼️ 图片格式": {
                "formats": ["png", "jpg", "jpeg", "bmp", "gif", "webp", "ico", "tiff", "svg", "pdf"],
                "available": PIL_AVAILABLE,
            },
            "📄 文档格式": {
                "formats": ["txt", "md", "html", "docx", "pdf", "rtf", "odt"],
                "available": True,
            },
            "📊 数据格式": {
                "formats": ["csv", "xlsx", "json", "xml", "yaml", "parquet", "tsv"],
                "available": PANDAS_AVAILABLE,
            },
            "🎵 音频格式": {
                "formats": ["mp3", "wav", "ogg", "flac", "m4a", "aac"],
                "available": AUDIO_AVAILABLE,
            },
            "🎬 视频格式": {
                "formats": ["mp4", "avi", "mkv", "mov", "flv", "wmv", "webm", "gif"],
                "available": VIDEO_AVAILABLE,
            },
            "📦 压缩格式": {
                "formats": ["zip", "tar", "tar.gz", "tar.bz2"],
                "available": True,
            },
            "📚 电子书": {
                "formats": ["epub", "mobi", "azw3", "txt", "html", "pdf"],
                "available": EPUB_AVAILABLE or PDF_AVAILABLE,
            },
            "🔐 编码转换": {
                "formats": ["base64", "hex", "md5", "sha256"],
                "available": True,
            },
            "🎨 特殊功能": {
                "formats": ["qrcode", "thumbnail", "watermark"],
                "available": PIL_AVAILABLE,
            },
        }

        self.init_ui()
        self.apply_theme()
        self.load_settings()

    def init_ui(self):
        """初始化UI"""
        self.setWindowTitle("🚀 终极文件格式转换器 v4.0 Pro (PyQt5)")
        self.setGeometry(100, 100, 1200, 800)
        self.setAcceptDrops(True)

        # 创建菜单栏
        self.create_menu_bar()

        # 创建工具栏
        self.create_tool_bar()

        # 创建中央部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        # 主布局
        main_layout = QHBoxLayout(central_widget)

        # 创建分割器
        splitter = QSplitter(Qt.Horizontal)

        # 左侧控制面板
        left_panel = self.create_left_panel()

        # 右侧信息面板
        right_panel = self.create_right_panel()

        splitter.addWidget(left_panel)
        splitter.addWidget(right_panel)
        splitter.setStretchFactor(0, 3)
        splitter.setStretchFactor(1, 2)

        main_layout.addWidget(splitter)

        # 状态栏
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        self.status_bar.showMessage("🟢 就绪")

    def create_menu_bar(self):
        """创建菜单栏"""
        menubar = self.menuBar()

        # 文件菜单
        file_menu = menubar.addMenu("文件")

        open_action = QAction("打开文件", self)
        open_action.setShortcut("Ctrl+O")
        open_action.triggered.connect(self.browse_files)
        file_menu.addAction(open_action)

        batch_action = QAction("批量打开", self)
        batch_action.setShortcut("Ctrl+Shift+O")
        batch_action.triggered.connect(self.browse_multiple_files)
        file_menu.addAction(batch_action)

        file_menu.addSeparator()

        exit_action = QAction("退出", self)
        exit_action.setShortcut("Ctrl+Q")
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)

        # 视图菜单
        view_menu = menubar.addMenu("视图")

        theme_action = QAction("切换主题", self)
        theme_action.setShortcut("Ctrl+T")
        theme_action.triggered.connect(self.toggle_theme)
        view_menu.addAction(theme_action)

        # 帮助菜单
        help_menu = menubar.addMenu("帮助")

        about_action = QAction("关于", self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)

        deps_action = QAction("依赖检查", self)
        deps_action.triggered.connect(self.check_dependencies)
        help_menu.addAction(deps_action)

    def create_tool_bar(self):
        """创建工具栏"""
        toolbar = QToolBar()
        toolbar.setMovable(False)
        toolbar.setIconSize(QSize(32, 32))
        self.addToolBar(toolbar)

        # 添加文件
        add_btn = QPushButton("📁 添加文件")
        add_btn.clicked.connect(self.browse_files)
        toolbar.addWidget(add_btn)

        # 清空列表
        clear_btn = QPushButton("🗑️ 清空")
        clear_btn.clicked.connect(self.clear_files)
        toolbar.addWidget(clear_btn)

        toolbar.addSeparator()

        # 开始转换
        convert_btn = QPushButton("🚀 开始转换")
        convert_btn.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                padding: 8px 16px;
                font-weight: bold;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        convert_btn.clicked.connect(self.start_conversion)
        toolbar.addWidget(convert_btn)

    def create_left_panel(self):
        """创建左侧控制面板"""
        panel = QWidget()
        layout = QVBoxLayout(panel)

        # 文件列表组
        file_group = QGroupBox("📁 文件列表")
        file_layout = QVBoxLayout()

        self.file_list = QListWidget()
        self.file_list.setAlternatingRowColors(True)
        file_layout.addWidget(self.file_list)

        # 文件操作按钮
        file_btn_layout = QHBoxLayout()
        add_btn = QPushButton("添加文件")
        add_btn.clicked.connect(self.browse_files)
        add_multiple_btn = QPushButton("批量添加")
        add_multiple_btn.clicked.connect(self.browse_multiple_files)
        remove_btn = QPushButton("移除选中")
        remove_btn.clicked.connect(self.remove_selected_files)

        file_btn_layout.addWidget(add_btn)
        file_btn_layout.addWidget(add_multiple_btn)
        file_btn_layout.addWidget(remove_btn)
        file_layout.addLayout(file_btn_layout)

        file_group.setLayout(file_layout)
        layout.addWidget(file_group)

        # 转换设置组
        settings_group = QGroupBox("⚙️ 转换设置")
        settings_layout = QVBoxLayout()

        # 类别选择
        category_label = QLabel("选择类别:")
        settings_layout.addWidget(category_label)

        self.category_buttons = QButtonGroup()
        category_grid = QGridLayout()

        categories = list(self.format_categories.keys())
        for i, cat in enumerate(categories):
            row, col = divmod(i, 3)
            info = self.format_categories[cat]
            rb = QRadioButton(cat)
            rb.setEnabled(info["available"])
            rb.toggled.connect(self.update_format_options)
            self.category_buttons.addButton(rb, i)
            category_grid.addWidget(rb, row, col)

        settings_layout.addLayout(category_grid)

        # 输出格式
        format_layout = QHBoxLayout()
        format_label = QLabel("输出格式:")
        self.format_combo = QComboBox()
        format_layout.addWidget(format_label)
        format_layout.addWidget(self.format_combo)
        settings_layout.addLayout(format_layout)

        # 质量设置
        quality_layout = QHBoxLayout()
        quality_label = QLabel("质量:")
        self.quality_slider = QSlider(Qt.Horizontal)
        self.quality_slider.setRange(1, 100)
        self.quality_slider.setValue(85)
        self.quality_value_label = QLabel("85")
        self.quality_slider.valueChanged.connect(
            lambda v: self.quality_value_label.setText(str(v))
        )
        quality_layout.addWidget(quality_label)
        quality_layout.addWidget(self.quality_slider)
        quality_layout.addWidget(self.quality_value_label)
        settings_layout.addLayout(quality_layout)

        # 其他选项
        self.keep_original_cb = QCheckBox("保留原文件")
        self.keep_original_cb.setChecked(True)
        self.auto_open_cb = QCheckBox("转换后自动打开")

        settings_layout.addWidget(self.keep_original_cb)
        settings_layout.addWidget(self.auto_open_cb)

        settings_group.setLayout(settings_layout)
        layout.addWidget(settings_group)

        # 输出位置组
        output_group = QGroupBox("📍 输出位置")
        output_layout = QVBoxLayout()

        self.output_label = QLabel("默认：与源文件相同位置")
        self.output_label.setWordWrap(True)
        output_layout.addWidget(self.output_label)

        browse_output_btn = QPushButton("选择输出文件夹")
        browse_output_btn.clicked.connect(self.browse_output_folder)
        output_layout.addWidget(browse_output_btn)

        output_group.setLayout(output_layout)
        layout.addWidget(output_group)

        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        layout.addWidget(self.progress_bar)

        return panel

    def create_right_panel(self):
        """创建右侧信息面板"""
        panel = QWidget()
        layout = QVBoxLayout(panel)

        # 标签页
        self.tab_widget = QTabWidget()

        # 日志标签页
        log_widget = QWidget()
        log_layout = QVBoxLayout(log_widget)
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setStyleSheet("QTextEdit { font-family: Consolas, monospace; }")
        log_layout.addWidget(self.log_text)
        self.tab_widget.addTab(log_widget, "📋 转换日志")

        # 文件信息标签页
        info_widget = QWidget()
        info_layout = QVBoxLayout(info_widget)
        self.info_text = QTextEdit()
        self.info_text.setReadOnly(True)
        self.info_text.setStyleSheet("QTextEdit { font-family: Consolas, monospace; }")
        info_layout.addWidget(self.info_text)
        self.tab_widget.addTab(info_widget, "ℹ️ 文件信息")

        # 历史记录标签页
        history_widget = QWidget()
        history_layout = QVBoxLayout(history_widget)
        self.history_table = QTableWidget()
        self.history_table.setColumnCount(5)
        self.history_table.setHorizontalHeaderLabels(["时间", "文件数", "成功", "失败", "用时"])
        self.history_table.horizontalHeader().setStretchLastSection(True)
        history_layout.addWidget(self.history_table)
        self.tab_widget.addTab(history_widget, "📜 历史记录")

        # 预览标签页
        preview_widget = QWidget()
        preview_layout = QVBoxLayout(preview_widget)
        self.preview_label = QLabel("选择图片文件查看预览")
        self.preview_label.setAlignment(Qt.AlignCenter)
        self.preview_label.setMinimumHeight(300)
        preview_layout.addWidget(self.preview_label)
        self.tab_widget.addTab(preview_widget, "👁️ 预览")

        # 帮助标签页
        help_widget = QWidget()
        help_layout = QVBoxLayout(help_widget)
        help_text = QTextEdit()
        help_text.setReadOnly(True)
        help_text.setPlainText(self.get_help_text())
        help_layout.addWidget(help_text)
        self.tab_widget.addTab(help_widget, "❓ 帮助")

        layout.addWidget(self.tab_widget)

        # 文件列表选择事件
        self.file_list.currentRowChanged.connect(self.on_file_selected)

        return panel

    # ========== 文件操作 ==========
    def browse_files(self):
        """选择单个文件"""
        files, _ = QFileDialog.getOpenFileNames(
            self,
            "选择要转换的文件",
            "",
            "所有文件 (*.*)"
        )
        if files:
            for file in files:
                if file not in self.input_files:
                    self.input_files.append(file)
                    self.file_list.addItem(Path(file).name)
            self.log_message(f"✅ 已添加 {len(files)} 个文件")

    def browse_multiple_files(self):
        """批量选择文件"""
        self.browse_files()

    def remove_selected_files(self):
        """移除选中的文件"""
        selected_items = self.file_list.selectedItems()
        if not selected_items:
            return

        for item in selected_items:
            row = self.file_list.row(item)
            self.file_list.takeItem(row)
            if row < len(self.input_files):
                self.input_files.pop(row)

        self.log_message("🗑️ 已移除选中文件")

    def clear_files(self):
        """清空文件列表"""
        self.input_files.clear()
        self.file_list.clear()
        self.info_text.clear()
        self.log_message("🗑️ 已清空文件列表")

    def browse_output_folder(self):
        """选择输出文件夹"""
        folder = QFileDialog.getExistingDirectory(self, "选择输出文件夹")
        if folder:
            self.output_folder = folder
            self.output_label.setText(folder)
            self.log_message(f"📂 输出位置: {folder}")

    def on_file_selected(self, row):
        """文件选中事件"""
        if row < 0 or row >= len(self.input_files):
            return

        file_path = self.input_files[row]
        self.show_file_info(file_path)
        self.show_file_preview(file_path)

    def show_file_info(self, filepath: str):
        """显示文件信息"""
        try:
            path = Path(filepath)
            stat = path.stat()

            info = f"""文件名: {path.name}
路径: {path.parent}
大小: {self.format_size(stat.st_size)}
创建时间: {datetime.fromtimestamp(stat.st_ctime).strftime('%Y-%m-%d %H:%M:%S')}
修改时间: {datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S')}
扩展名: {path.suffix}
MD5: {self.calculate_md5(filepath)[:32]}
"""
            self.info_text.setPlainText(info)

        except Exception as e:
            self.info_text.setPlainText(f"无法获取文件信息: {e}")

    def show_file_preview(self, filepath: str):
        """显示文件预览"""
        try:
            ext = Path(filepath).suffix.lower()[1:]

            if ext in ["png", "jpg", "jpeg", "bmp", "gif", "webp"] and PIL_AVAILABLE:
                pixmap = QPixmap(filepath)
                if not pixmap.isNull():
                    scaled_pixmap = pixmap.scaled(
                        400, 400,
                        Qt.KeepAspectRatio,
                        Qt.SmoothTransformation
                    )
                    self.preview_label.setPixmap(scaled_pixmap)
                else:
                    self.preview_label.setText("无法加载图片")
            elif ext in ["txt", "md", "log"]:
                with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read(1000)
                self.preview_label.setText(content + "\n..." if len(content) >= 1000 else content)
            else:
                self.preview_label.setText("不支持预览此文件类型")

        except Exception as e:
            self.preview_label.setText(f"预览失败: {e}")

    # ========== 转换相关 ==========
    def update_format_options(self):
        """更新格式选项"""
        checked_button = self.category_buttons.checkedButton()
        if not checked_button:
            return

        category = checked_button.text()
        if category in self.format_categories:
            info = self.format_categories[category]
            if info["available"]:
                self.format_combo.clear()
                self.format_combo.addItems(info["formats"])
                self.log_message(f"📋 已加载 {len(info['formats'])} 个格式选项")
            else:
                self.format_combo.clear()
                self.log_message(f"⚠️ 类别 {category} 依赖未安装", "warning")

    def start_conversion(self):
        """开始转换"""
        if not self.input_files:
            QMessageBox.warning(self, "警告", "请先添加要转换的文件")
            return

        checked_button = self.category_buttons.checkedButton()
        if not checked_button:
            QMessageBox.warning(self, "警告", "请选择转换类别")
            return

        if not self.format_combo.currentText():
            QMessageBox.warning(self, "警告", "请选择输出格式")
            return

        # 收集设置
        settings = {
            "category": checked_button.text(),
            "output_format": self.format_combo.currentText(),
            "quality": self.quality_slider.value(),
            "keep_original": self.keep_original_cb.isChecked(),
            "auto_open": self.auto_open_cb.isChecked(),
            "output_folder": self.output_folder,
        }

        # 创建并启动工作线程
        self.worker = ConversionWorker(self.engine, self.input_files, settings)
        self.worker.progress.connect(self.on_conversion_progress)
        self.worker.finished.connect(self.on_conversion_finished)
        self.worker.log.connect(self.log_message)

        self.progress_bar.setValue(0)
        self.status_bar.showMessage("🔄 转换中...")
        self.worker.start()

    def on_conversion_progress(self, current, total, message):
        """转换进度更新"""
        progress = int((current / total) * 100)
        self.progress_bar.setValue(progress)
        self.status_bar.showMessage(message)

    def on_conversion_finished(self, success, failed, elapsed):
        """转换完成"""
        self.progress_bar.setValue(100)
        self.status_bar.showMessage(f"✅ 完成: {success}/{success + failed}, 失败: {failed}")

        # 添加到历史
        self.add_to_history(success, failed, elapsed)

        # 显示结果
        total = success + failed
        msg = f"转换完成！\n✅ 成功: {success}\n❌ 失败: {failed}\n总计: {total}\n⏱️ 用时: {elapsed:.2f} 秒"

        if success > 0:
            QMessageBox.information(self, "完成", msg)
        else:
            QMessageBox.critical(self, "错误", msg)

    def add_to_history(self, success, failed, elapsed):
        """添加到历史记录"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        total = success + failed

        row = self.history_table.rowCount()
        self.history_table.insertRow(row)

        self.history_table.setItem(row, 0, QTableWidgetItem(timestamp))
        self.history_table.setItem(row, 1, QTableWidgetItem(str(total)))
        self.history_table.setItem(row, 2, QTableWidgetItem(str(success)))
        self.history_table.setItem(row, 3, QTableWidgetItem(str(failed)))
        self.history_table.setItem(row, 4, QTableWidgetItem(f"{elapsed:.2f}s"))

    # ========== 日志相关 ==========
    def log_message(self, message: str, level: str = "info"):
        """添加日志"""
        timestamp = datetime.now().strftime("%H:%M:%S")

        color_map = {
            "info": "#000000",
            "success": "#27ae60",
            "warning": "#f39c12",
            "error": "#e74c3c",
        }

        color = color_map.get(level, "#000000")
        self.log_text.append(f'<span style="color: {color};">[{timestamp}] {message}</span>')

    # ========== 工具函数 ==========
    @staticmethod
    def format_size(bytes_size: float) -> str:
        """格式化文件大小"""
        for unit in ["B", "KB", "MB", "GB"]:
            if bytes_size < 1024:
                return f"{bytes_size:.2f} {unit}"
            bytes_size /= 1024
        return f"{bytes_size:.2f} TB"

    @staticmethod
    def calculate_md5(filename: str) -> str:
        """计算MD5"""
        try:
            hash_md5 = hashlib.md5()
            with open(filename, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except:
            return "无法计算"

    def get_help_text(self) -> str:
        """获取帮助文本"""
        return """
╔═══════════════════════════════════════════╗
║    终极文件格式转换器 v4.0 Pro (PyQt5)    ║
╚═══════════════════════════════════════════╝

【新增功能】
✨ 全新 PyQt5 界面 - 现代化设计
✨ 拖拽文件支持 - 方便快捷
✨ 文件预览功能 - 实时查看
✨ 主题切换 - 亮色/暗色模式
✨ 详细进度显示 - 实时反馈

【支持的格式】
📸 图片: PNG, JPG, BMP, GIF, WEBP, SVG, PDF
📝 文档: TXT, MD, HTML, DOCX, PDF
📊 数据: CSV, XLSX, JSON, XML, YAML
🎵 音频: MP3, WAV, OGG, FLAC, M4A
🎬 视频: MP4, AVI, MKV, MOV, WEBM, GIF
📦 压缩: ZIP, TAR, TAR.GZ
📚 电子书: EPUB, MOBI, AZW3, PDF
🔐 编码: Base64, Hex, MD5, SHA256
🎨 特殊: 二维码, 缩略图, 水印

【快捷键】
Ctrl+O : 打开文件
Ctrl+Shift+O : 批量打开
Ctrl+T : 切换主题
Ctrl+Q : 退出程序

【推荐依赖】
pip install PyQt5 Pillow pandas openpyxl python-docx \\
    beautifulsoup4 pypdf reportlab pydub moviepy \\
    ebooklib pyyaml qrcode cairosvg

【注意事项】
- 音视频转换需要安装 FFmpeg
- PDF中文需要系统中文字体
- MOBI/AZW3 需要安装 Calibre
"""

    # ========== 主题相关 ==========
    def toggle_theme(self):
        """切换主题"""
        self.current_theme = "dark" if self.current_theme == "light" else "light"
        self.apply_theme()
        self.log_message(f"🎨 已切换到 {self.current_theme} 主题")

    def apply_theme(self):
        """应用主题"""
        if self.current_theme == "dark":
            self.setStyleSheet("""
                QMainWindow, QWidget {
                    background-color: #2b2b2b;
                    color: #ffffff;
                }
                QGroupBox {
                    border: 1px solid #444;
                    border-radius: 5px;
                    margin-top: 10px;
                    padding-top: 10px;
                    font-weight: bold;
                }
                QGroupBox::title {
                    subcontrol-origin: margin;
                    left: 10px;
                    padding: 0 5px;
                }
                QPushButton {
                    background-color: #3a3a3a;
                    color: white;
                    border: 1px solid #555;
                    padding: 6px 12px;
                    border-radius: 4px;
                }
                QPushButton:hover {
                    background-color: #4a4a4a;
                }
                QListWidget, QTextEdit, QComboBox {
                    background-color: #1e1e1e;
                    color: #ffffff;
                    border: 1px solid #444;
                }
                QProgressBar {
                    border: 1px solid #444;
                    border-radius: 4px;
                    text-align: center;
                }
                QProgressBar::chunk {
                    background-color: #27ae60;
                }
            """)
        else:
            self.setStyleSheet("""
                QGroupBox {
                    border: 1px solid #ddd;
                    border-radius: 5px;
                    margin-top: 10px;
                    padding-top: 10px;
                    font-weight: bold;
                }
                QGroupBox::title {
                    subcontrol-origin: margin;
                    left: 10px;
                    padding: 0 5px;
                }
                QPushButton {
                    background-color: #f0f0f0;
                    border: 1px solid #ccc;
                    padding: 6px 12px;
                    border-radius: 4px;
                }
                QPushButton:hover {
                    background-color: #e0e0e0;
                }
                QProgressBar {
                    border: 1px solid #ccc;
                    border-radius: 4px;
                    text-align: center;
                }
                QProgressBar::chunk {
                    background-color: #27ae60;
                }
            """)

    # ========== 拖拽支持 ==========
    def dragEnterEvent(self, event: QDragEnterEvent):
        """拖拽进入事件"""
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dropEvent(self, event: QDropEvent):
        """放置事件"""
        urls = event.mimeData().urls()
        for url in urls:
            file_path = url.toLocalFile()
            if os.path.isfile(file_path) and file_path not in self.input_files:
                self.input_files.append(file_path)
                self.file_list.addItem(Path(file_path).name)

        self.log_message(f"✅ 已拖拽添加 {len(urls)} 个文件")

    # ========== 其他功能 ==========
    def check_dependencies(self):
        """检查依赖"""
        missing = []
        if not PIL_AVAILABLE:
            missing.append("Pillow")
        if not PANDAS_AVAILABLE:
            missing.append("pandas + openpyxl")
        if not DOC_AVAILABLE:
            missing.append("python-docx + beautifulsoup4")
        if not AUDIO_AVAILABLE:
            missing.append("pydub + ffmpeg")
        if not VIDEO_AVAILABLE:
            missing.append("moviepy + ffmpeg")
        if not PDF_AVAILABLE:
            missing.append("pypdf + reportlab")
        if not EPUB_AVAILABLE:
            missing.append("ebooklib")
        if not YAML_AVAILABLE:
            missing.append("pyyaml")
        if not QRCODE_AVAILABLE:
            missing.append("qrcode")
        if not CAIROSVG_AVAILABLE:
            missing.append("cairosvg (SVG支持)")
        if not CALIBRE_AVAILABLE:
            missing.append("Calibre (电子书增强)")

        if missing:
            msg = "缺少以下依赖库:\n\n" + "\n".join([f"• {dep}" for dep in missing])
            QMessageBox.warning(self, "依赖检查", msg)
        else:
            QMessageBox.information(self, "依赖检查", "✅ 所有依赖库已安装!")

    def show_about(self):
        """显示关于"""
        QMessageBox.about(
            self,
            "关于",
            """<h2>终极文件格式转换器 v4.0 Pro</h2>
            <p>基于 PyQt5 的现代化文件转换工具</p>
            <p><b>特性:</b></p>
            <ul>
                <li>支持多种文件格式互转</li>
                <li>批量转换处理</li>
                <li>拖拽文件支持</li>
                <li>实时预览功能</li>
                <li>主题切换</li>
            </ul>
            <p><b>作者:</b> Ultimate Converter Team</p>
            <p><b>版本:</b> 4.0 Pro (PyQt5)</p>
            """
        )

    def load_settings(self):
        """加载设置"""
        try:
            quality = self.settings.value("quality", 85, type=int)
            self.quality_slider.setValue(quality)

            keep_original = self.settings.value("keep_original", True, type=bool)
            self.keep_original_cb.setChecked(keep_original)

            theme = self.settings.value("theme", "light")
            self.current_theme = theme
            self.apply_theme()
        except:
            pass

    def save_settings(self):
        """保存设置"""
        try:
            self.settings.setValue("quality", self.quality_slider.value())
            self.settings.setValue("keep_original", self.keep_original_cb.isChecked())
            self.settings.setValue("theme", self.current_theme)
        except:
            pass

    def closeEvent(self, event):
        """关闭事件"""
        self.save_settings()
        event.accept()


# ================== 主函数 ==================
def main():
    app = QApplication(sys.argv)
    app.setApplicationName("终极文件格式转换器")
    app.setStyle("Fusion")

    # 设置应用图标（如果有）
    # app.setWindowIcon(QIcon("icon.png"))

    window = UltimateFileConverter()
    window.show()

    sys.exit(app.exec_())


if __name__ == "__main__":
    main()